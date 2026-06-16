# -*- coding: utf-8 -*-
"""
Сглобява Masters_Diplomna_Arkan_Ahmedov.docx от главите в chapters/.
Оформление по изискванията: Times New Roman, тяло 14pt двустранно,
H1 20 / H2 16 / H3 14 (bold), корица с лого, автоматично съдържание,
номерация на страниците долу вдясно.
"""
import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FONT = "Times New Roman"

# Главите се сглобяват в този ред. Допълва се с напредването на писането.
CHAPTERS = [
    "chapters/00-vavedenie.md",
    "chapters/01-glava1.md",
    "chapters/02-glava2.md",
    "chapters/03-glava3.md",
    "chapters/04-glava4.md",
    "chapters/05-glava5.md",
    "chapters/06-glava6.md",
    "chapters/07-glava7.md",
    "chapters/08-glava8.md",
    "chapters/09-glava9.md",
    "chapters/10-zaklyuchenie.md",
    "chapters/11-prilozhenia.md",
    "chapters/12-deklaracia.md",
    "chapters/13-literatura.md",
]


# ---------- ниско ниво помощници ----------
def set_cell_font(run, size=14, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)


def style_base_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
    if rfonts.getparent() is None:
        rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)

    for name, size, sb, sa in (("Heading 1", 20, 18, 12),
                               ("Heading 2", 16, 12, 6),
                               ("Heading 3", 14, 6, 4)):
        st = doc.styles[name]
        st.font.name = FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
        if rfonts.getparent() is None:
            rpr.append(rfonts)
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rfonts.set(qn(a), FONT)
        st.paragraph_format.space_before = Pt(sb)
        st.paragraph_format.space_after = Pt(sa)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.keep_with_next = True


def add_page_number_footer(section):
    """Номер на страницата долу вдясно, чрез поле PAGE."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    for t, val in (('begin', None), ('instr', 'PAGE'), ('end', None)):
        el = OxmlElement('w:fldChar') if t != 'instr' else OxmlElement('w:instrText')
        if t == 'instr':
            el.set(qn('xml:space'), 'preserve'); el.text = ' PAGE '
        else:
            el.set(qn('w:fldCharType'), t)
        run._element.append(el)
    set_cell_font(run, size=11)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fb = OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = 'TOC \\o "1-3" \\h \\z \\u'
    fs = OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'), 'separate')
    ph = OxmlElement('w:r'); pht = OxmlElement('w:t')
    pht.text = 'Съдържанието се обновява в Word с десен бутон, Update Field, или F9.'
    ph.append(pht)
    fe = OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'), 'end')
    for el in (fb, it, fs, ph, fe):
        run._element.append(el)


def centered(doc, text, size=14, bold=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    set_cell_font(r, size=size, bold=bold)
    return p


# ---------- корица ----------
def build_cover(doc):
    logo = os.path.join(HERE, "assets", "tu-sofia-logo.jpeg")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    if os.path.exists(logo):
        p.add_run().add_picture(logo, width=Cm(3.0))

    centered(doc, "ТЕХНИЧЕСКИ УНИВЕРСИТЕТ СОФИЯ", size=16, bold=True, space_before=10, space_after=2)
    centered(doc, "Факултет по приложна математика и информатика", size=14, space_after=2)
    centered(doc, "Катедра „Информатика“", size=14, space_after=24)

    centered(doc, "ДИПЛОМНА РАБОТА", size=22, bold=True, space_before=18, space_after=4)
    centered(doc, "за придобиване на образователно-квалификационна степен „Магистър“", size=13, space_after=28)

    centered(doc, "на тема:", size=14, space_after=8)
    centered(doc,
             "„Платформа за управление и агрегиране на вградими уеб компоненти "
             "за разработване на уеб приложения за туризма“",
             size=16, bold=True, space_after=36)

    # данни
    for label, value in (
        ("Дипломант:", "Аркан Ахмедов Ахмедов"),
        ("Факултетен номер:", "791325012"),
        ("Специалност:", "Информатика и софтуерни науки"),
        ("Образователно-квалификационна степен:", "Магистър"),
        ("Научен ръководител:", "проф. Малинка Иванова"),
    ):
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_after = Pt(2)
        r1 = pp.add_run(label + " "); set_cell_font(r1, size=14, bold=True)
        r2 = pp.add_run(value); set_cell_font(r2, size=14)

    centered(doc, "София, 2026", size=14, bold=True, space_before=40)


# ---------- markdown -> параграфи ----------
def add_runs_with_bold(paragraph, text):
    # минимален inline parser за **bold**
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); set_cell_font(r, size=14, bold=True)
        else:
            r = paragraph.add_run(part); set_cell_font(r, size=14)


def _is_table_sep(line):
    return bool(re.match(r'^\s*\|?[\s:\-\|]+\|?\s*$', line)) and '-' in line


def _table_cells(line):
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def make_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    table.alignment = 1  # center
    for j, h in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.paragraphs[0].text = ''
        r = cell.paragraphs[0].add_run(h)
        set_cell_font(r, size=12, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for j, val in enumerate(row):
            if j >= len(cells):
                break
            cells[j].paragraphs[0].text = ''
            r = cells[j].paragraphs[0].add_run(val)
            set_cell_font(r, size=12)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def png_size(path):
    import struct
    with open(path, 'rb') as f:
        head = f.read(24)
    if head[:8] != b'\x89PNG\r\n\x1a\n':
        return None
    w, h = struct.unpack('>II', head[16:24])
    return w, h


def add_figure(doc, path, caption):
    full = path if os.path.isabs(path) else os.path.join(HERE, path)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    dims = png_size(full)
    width = Cm(14)
    if dims:
        w, h = dims
        ratio = h / w
        target_w = Cm(14)
        if Cm(14) * ratio > Cm(19):
            target_w = int(Cm(19) / ratio)
        width = target_w
    if os.path.exists(full):
        p.add_run().add_picture(full, width=width)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    r = cap.add_run(caption); set_cell_font(r, size=12); r.italic = True


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F4F3')
    ppr.append(shd)
    for idx, ln in enumerate(code_lines):
        if idx > 0:
            p.add_run().add_break()
        r = p.add_run(ln.replace('\t', '    '))
        r.font.name = 'Consolas'; r.font.size = Pt(10.5)
        rpr = r._element.get_or_add_rPr(); rf = OxmlElement('w:rFonts')
        for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rf.set(qn(a), 'Consolas')
        rpr.append(rf)


def render_markdown(doc, md):
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1; continue

        # кодов блок: ```...```
        if line.lstrip().startswith('```'):
            i += 1
            code = []
            while i < len(lines) and not lines[i].lstrip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1  # затварящ ```
            add_code_block(doc, code)
            continue

        # изображение/фигура: ![Надпис](път)
        m_img = re.match(r'^!\[(.*?)\]\((.*?)\)\s*$', line)
        if m_img:
            add_figure(doc, m_img.group(2).strip(), m_img.group(1).strip())
            i += 1; continue

        # таблица
        if line.lstrip().startswith('|'):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                block.append(lines[i]); i += 1
            header = _table_cells(block[0])
            data = [_table_cells(b) for b in block[1:] if not _is_table_sep(b)]
            make_table(doc, header, data)
            continue

        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r'^(Таблица|Фигура)\s', line):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(line); set_cell_font(r, size=12); r.italic = True
        elif re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(0.6)
            add_runs_with_bold(p, line)
        elif line.startswith('> '):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip()); set_cell_font(r, size=12); r.italic = True
        else:
            p = doc.add_paragraph()
            add_runs_with_bold(p, line)
        i += 1


# ---------- сглобяване ----------
def build():
    doc = Document()
    for s in doc.sections:
        s.page_height = Cm(29.7); s.page_width = Cm(21.0)
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.5)
    style_base_styles(doc)

    # Корица (без номер на страницата)
    build_cover(doc)

    # Нова секция за тялото, с номерация долу вдясно
    doc.add_section(WD_SECTION.NEW_PAGE)
    body = doc.sections[-1]
    body.different_first_page_header_footer = False
    add_page_number_footer(body)

    # Съдържание
    doc.add_heading("СЪДЪРЖАНИЕ", level=1)
    add_toc(doc)
    doc.add_page_break()

    # Глави
    for rel in CHAPTERS:
        path = os.path.join(HERE, rel)
        with open(path, encoding='utf-8') as f:
            render_markdown(doc, f.read())
        doc.add_page_break()

    out = os.path.join(HERE, "Masters_Diplomna_Arkan_Ahmedov.docx")
    doc.save(out)
    print("Записан:", out)


if __name__ == "__main__":
    build()
